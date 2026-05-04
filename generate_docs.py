"""Generates Styleus_Architecture.docx on the Desktop."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)

# ── Colour palette ───────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1e, 0x21, 0x60)
BLUE   = RGBColor(0x3b, 0x82, 0xf6)
GREEN  = RGBColor(0x16, 0xa3, 0x4a)
PURPLE = RGBColor(0x7c, 0x3a, 0xed)
ORANGE = RGBColor(0xea, 0x58, 0x0c)
GREY   = RGBColor(0x6b, 0x72, 0x80)
WHITE  = RGBColor(0xff, 0xff, 0xff)
BLACK  = RGBColor(0x11, 0x18, 0x27)
CYAN   = RGBColor(0x06, 0xb6, 0xd4)
RED    = RGBColor(0xdc, 0x26, 0x26)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.size  = Pt(22)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    return p

def h2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(14)
    run.font.bold  = True
    run.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    return p

def h3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    return p

def body(text, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(3)
    return p

def mono(text, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(1)
    return p

def bullet(text, level=0, color=None):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    if color:
        run.font.color.rgb = color
    p.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after = Pt(2)
    return p

def rule():
    p = doc.add_paragraph()
    run = p.add_run('─' * 95)
    run.font.color.rgb = RGBColor(0xd1, 0xd5, 0xdb)
    run.font.size = Pt(8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)

def header_table(cols, widths, header_color='1e2160'):
    t = doc.add_table(rows=1, cols=len(cols))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (col, w) in enumerate(zip(cols, widths)):
        cell = t.rows[0].cells[i]
        cell.width = Inches(w)
        set_cell_bg(cell, header_color)
        p = cell.paragraphs[0]
        run = p.add_run(col)
        run.font.bold  = True
        run.font.color.rgb = WHITE
        run.font.size  = Pt(9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return t

def add_row(table, values, widths, alt=False):
    row = table.add_row()
    for i, (val, w) in enumerate(zip(values, widths)):
        cell = row.cells[i]
        cell.width = Inches(w)
        if alt:
            set_cell_bg(cell, 'f3f4f6')
        p = cell.paragraphs[0]
        run = p.add_run(str(val))
        run.font.size = Pt(9)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return row

# ═════════════════════════════════════════════════════════════════════════════
# COVER
# ═════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('STYLEUS')
run.font.size  = Pt(36)
run.font.bold  = True
run.font.color.rgb = NAVY
p.paragraph_format.space_before = Pt(24)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Application Architecture & Module Reference')
run.font.size  = Pt(16)
run.font.color.rgb = BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Flask · PostgreSQL · Docker · Gunicorn · Nginx')
run.font.size  = Pt(11)
run.font.color.rgb = GREY

doc.add_paragraph()
rule()
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ═════════════════════════════════════════════════════════════════════════════
h1('1.  System Overview')
body(
    'Styleus is a full-stack e-commerce platform built with Python/Flask. '
    'It supports physical and digital product sales, a peer-to-peer used-item marketplace, '
    'IT services listings, tech enquiries, Stripe payments, and a full admin panel.'
)

doc.add_paragraph()
h2('High-Level Architecture')
mono('┌─────────────────────────────────────────────────────────────────────┐')
mono('│                        INTERNET / BROWSER                           │')
mono('└────────────────────────────┬────────────────────────────────────────┘')
mono('                             │  HTTPS :443')
mono('                             ▼')
mono('┌─────────────────────────────────────────────────────────────────────┐')
mono('│                          NGINX (Reverse Proxy)                      │')
mono('│               Handles SSL, static files, rate limiting              │')
mono('└────────────────────────────┬────────────────────────────────────────┘')
mono('                             │  HTTP :5000')
mono('                             ▼')
mono('┌─────────────────────────────────────────────────────────────────────┐')
mono('│                     GUNICORN (WSGI Server)                          │')
mono('│              2 sync workers  ·  container: stylus_app               │')
mono('└────────────────────────────┬────────────────────────────────────────┘')
mono('                             │')
mono('                             ▼')
mono('┌─────────────────────────────────────────────────────────────────────┐')
mono('│                      FLASK APPLICATION                              │')
mono('│  app.py  ·  config.py  ·  models.py  ·  routes/  ·  templates/     │')
mono('└──────────┬────────────────────────────────────────┬─────────────────┘')
mono('           │ SQLAlchemy ORM                          │ static files')
mono('           ▼                                         ▼')
mono('┌──────────────────────┐              ┌──────────────────────────────┐')
mono('│  PostgreSQL :5432    │              │  /static/uploads/  (volume)  │')
mono('│  container:stylus_db │              │  /static/js,css,images       │')
mono('└──────────────────────┘              └──────────────────────────────┘')
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 2. PROJECT FILE STRUCTURE
# ═════════════════════════════════════════════════════════════════════════════
h1('2.  Project File Structure')
mono('styleus/')
mono('├── app.py                ← Flask app factory, blueprints, CLI seed, error handlers')
mono('├── config.py             ← Config class (env vars, upload paths, Stripe keys)')
mono('├── models.py             ← All SQLAlchemy models (10 tables)')
mono('├── Dockerfile            ← Python 3.11-slim, gunicorn, entrypoint')
mono('├── entrypoint.sh         ← db.create_all() + flask seed + gunicorn start')
mono('├── requirements.txt      ← Python dependencies')
mono('│')
mono('├── routes/')
mono('│   ├── shop.py           ← Public storefront  (prefix: /)')
mono('│   ├── auth.py           ← Login / Register / Logout  (prefix: /auth)')
mono('│   ├── cart.py           ← Cart  AJAX API  (prefix: /cart)')
mono('│   ├── checkout.py       ← Checkout + Stripe  (prefix: /checkout)')
mono('│   ├── orders.py         ← Order history  (prefix: /orders)')
mono('│   ├── sell.py           ← User listings (sell used items)  (prefix: /)')
mono('│   └── admin.py          ← Admin CRUD panel  (prefix: /admin)')
mono('│')
mono('├── templates/')
mono('│   ├── base.html         ← Public layout (navbar, footer, flash msgs)')
mono('│   ├── index.html        ← Homepage')
mono('│   ├── services.html     ← Services listing')
mono('│   ├── enquiry.html      ← Tech Enquiries form')
mono('│   ├── _product_card.html← Reusable product card partial')
mono('│   ├── auth/             ← login.html, register.html')
mono('│   ├── shop/             ← products.html, product.html, search_results.html')
mono('│   ├── cart/             ← cart.html')
mono('│   ├── checkout/         ← checkout.html, success.html')
mono('│   ├── orders/           ← history.html, detail.html')
mono('│   ├── sell/             ← sell_form.html, my_listings.html, listings.html')
mono('│   └── admin/            ← base.html + dashboard, products, orders, users,')
mono('│                           categories, listings, services, sell_categories,')
mono('│                           enquiries, product_form, service_form, order_detail')
mono('│')
mono('└── static/')
mono('    ├── js/cart.js        ← Cart AJAX logic')
mono('    ├── css/              ← Stylesheets')
mono('    ├── uploads/          ← User-uploaded product & listing images (volume mount)')
mono('    └── images/           ← Static site images')
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 3. BLUEPRINTS & ROUTES
# ═════════════════════════════════════════════════════════════════════════════
h1('3.  Blueprints & Routes')
body('Flask organises routes into 7 blueprints, each registered in app.py.')
doc.add_paragraph()

blueprints = [
    ('shop_bp',     '/shop.py',    '/',         'Public storefront',       'GREEN'),
    ('auth_bp',     '/auth.py',    '/auth',     'Authentication',          'BLUE'),
    ('cart_bp',     '/cart.py',    '/cart',     'Cart (AJAX JSON API)',    'ORANGE'),
    ('checkout_bp', '/checkout.py','/checkout', 'Checkout & Stripe',       'PURPLE'),
    ('orders_bp',   '/orders.py',  '/orders',   'Order history',           'CYAN'),
    ('sell_bp',     '/sell.py',    '/',         'User used-item listings', 'ORANGE'),
    ('admin_bp',    '/admin.py',   '/admin',    'Admin panel (protected)', 'RED'),
]

cols   = ['Blueprint', 'File', 'URL Prefix', 'Purpose']
widths = [1.4, 1.5, 1.2, 3.1]
t = header_table(cols, widths)
for i, (bp, f, prefix, purpose, _) in enumerate(blueprints):
    add_row(t, [bp, 'routes' + f, prefix, purpose], widths, alt=(i % 2 == 0))

doc.add_paragraph()
rule()

# ── Route detail tables ───────────────────────────────────────────────────────
route_data = {
    'shop_bp — Public Storefront  (routes/shop.py)': [
        ('GET',  '/',                    'Homepage — featured products, services, hero'),
        ('GET',  '/products',            'Product catalogue with search, filter, sort, pagination'),
        ('GET',  '/products/<slug>',     'Product detail page with images, reviews, related products'),
        ('POST', '/products/<slug>/review', 'Submit / update a star review (login required)'),
        ('POST', '/products/<int>/like', 'Toggle product like / unlike (AJAX, login required)'),
        ('GET',  '/services',            'Public IT services listing'),
        ('GET',  '/search',              'Combined search: shop products + used listings'),
        ('GET',  '/enquiry',             'Tech Enquiries form'),
        ('POST', '/enquiry',             'Submit enquiry (validate email, phone, details ≥25 chars)'),
    ],
    'auth_bp — Authentication  (routes/auth.py)': [
        ('GET',  '/auth/login',    'Login page'),
        ('POST', '/auth/login',    'Validate credentials; merge session cart on login'),
        ('GET',  '/auth/register', 'Registration page'),
        ('POST', '/auth/register', 'Create user account (pending admin approval)'),
        ('GET',  '/auth/logout',   'Log out current user'),
    ],
    'cart_bp — Cart API  (routes/cart.py)': [
        ('GET',  '/cart/',         'View cart (DB cart for logged-in, session cart for guests)'),
        ('POST', '/cart/add',      'Add item to cart (JSON) → returns updated cart_count'),
        ('POST', '/cart/update',   'Update item quantity (JSON) → returns total, subtotal'),
        ('POST', '/cart/remove',   'Remove item from cart (JSON) → returns total'),
        ('POST', '/cart/clear',    'Empty the entire cart'),
    ],
    'checkout_bp — Checkout  (routes/checkout.py)': [
        ('GET',  '/checkout/', 'Checkout form (currently disabled via flag)'),
        ('POST', '/checkout/', 'Process order → Stripe PaymentIntent → create Order + OrderItems'),
        ('GET',  '/checkout/success', 'Order confirmation page'),
    ],
    'orders_bp — Orders  (routes/orders.py)': [
        ('GET', '/orders/',         'Order history list (login required)'),
        ('GET', '/orders/<int:id>', 'Order detail (login required; admin sees all)'),
    ],
    'sell_bp — Used Item Marketplace  (routes/sell.py)': [
        ('GET',  '/sell',              'Submit listing form (login required)'),
        ('POST', '/sell',              'Create UserListing (pending admin approval)'),
        ('GET',  '/sell/my-listings',  'User\'s own listings with status'),
        ('GET',  '/sell/listings',     'Public approved listings'),
    ],
    'admin_bp — Admin Panel  (routes/admin.py)': [
        ('GET',  '/admin/',                          'Dashboard — stats, recent orders, pending alerts'),
        ('GET',  '/admin/products',                  'List all products (search, paginate)'),
        ('GET/POST', '/admin/products/new',          'Create product (with up to 4 image uploads)'),
        ('GET/POST', '/admin/products/<id>/edit',    'Edit product'),
        ('POST', '/admin/products/<id>/delete',      'Delete product'),
        ('POST', '/admin/products/<id>/toggle',      'Activate / deactivate product'),
        ('GET',  '/admin/orders',                    'List orders (filter by status)'),
        ('GET',  '/admin/orders/<id>',               'Order detail view'),
        ('POST', '/admin/orders/<id>/status',        'Update order status'),
        ('GET',  '/admin/users',                     'List users (search, pending filter)'),
        ('POST', '/admin/users/<id>/approve',        'Approve pending user registration'),
        ('POST', '/admin/users/<id>/toggle-admin',   'Grant / revoke admin role'),
        ('GET',  '/admin/categories',                'Manage product categories'),
        ('POST', '/admin/categories/new',            'Create category'),
        ('POST', '/admin/categories/<id>/edit',      'Edit category'),
        ('POST', '/admin/categories/<id>/delete',    'Delete category (unlinks products)'),
        ('GET',  '/admin/listings',                  'User listings (filter by status)'),
        ('POST', '/admin/listings/<id>/approve',     'Approve listing'),
        ('POST', '/admin/listings/<id>/reject',      'Reject listing'),
        ('GET/POST', '/admin/services',              'Manage IT services'),
        ('GET/POST', '/admin/services/new',          'Create service'),
        ('GET/POST', '/admin/services/<id>/edit',    'Edit service'),
        ('POST', '/admin/services/<id>/delete',      'Delete service'),
        ('POST', '/admin/services/<id>/toggle',      'Activate / deactivate service'),
        ('GET',  '/admin/sell-categories',           'Manage used-item sell categories'),
        ('GET',  '/admin/enquiries',                 'View all tech enquiries'),
        ('POST', '/admin/enquiries/<id>/delete',     'Delete an enquiry'),
        ('POST', '/admin/reviews/<id>/delete',       'Delete a product review'),
    ],
}

for section_title, rows in route_data.items():
    h3(section_title)
    cols   = ['Method', 'URL', 'Description']
    widths = [0.9, 2.6, 3.7]
    t = header_table(cols, widths, header_color='3235a4')
    for i, (method, url, desc) in enumerate(rows):
        add_row(t, [method, url, desc], widths, alt=(i % 2 == 0))
    doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 4. DATABASE MODELS
# ═════════════════════════════════════════════════════════════════════════════
h1('4.  Database Models  (models.py)')
body('All models live in models.py and are managed by SQLAlchemy. No Alembic — db.create_all() handles table creation on container start.')
doc.add_paragraph()

h2('Entity-Relationship Overview')
mono('┌───────────────┐       ┌─────────────────┐       ┌───────────────┐')
mono('│     User      │──┐    │    Product      │──┐    │   Category    │')
mono('│───────────────│  │    │─────────────────│  │    │───────────────│')
mono('│ id            │  │    │ id              │  └───▶│ id            │')
mono('│ name          │  │    │ name            │       │ name          │')
mono('│ email (uniq)  │  │    │ slug (uniq)     │       │ slug (uniq)   │')
mono('│ phone         │  │    │ description     │       │ description   │')
mono('│ password_hash │  │    │ price           │       └───────────────┘')
mono('│ is_admin      │  │    │ stock           │')
mono('│ is_approved   │  │    │ product_type    │  ┌──────────────────────┐')
mono('│ created_at    │  │    │ image_url 1-4   │  │      Review          │')
mono('└───────────────┘  │    │ is_active       │  │──────────────────────│')
mono('       │            │    │ is_featured     │◀─│ id                   │')
mono('       │ 1:1        │    │ created_at      │  │ product_id (FK)      │')
mono('       ▼            │    └─────────────────┘  │ user_id (FK)         │')
mono('┌───────────────┐  │             │             │ rating (1-5)         │')
mono('│     Cart      │  │             │             │ body                 │')
mono('│───────────────│  │      ┌──────┴───────┐     │ created_at           │')
mono('│ id            │  │      │  CartItem    │     │ UNIQUE(user,product) │')
mono('│ user_id (FK)  │  └─────▶│──────────────│     └──────────────────────┘')
mono('│ created_at    │         │ cart_id (FK) │')
mono('└───────────────┘         │ product_id   │  ┌──────────────────────┐')
mono('       │ 1:N              │ quantity     │  │   ProductLike        │')
mono('       ▼                  └──────────────┘  │──────────────────────│')
mono('┌───────────────┐                           │ user_id (FK)         │')
mono('│     Order     │                           │ product_id (FK)      │')
mono('│───────────────│                           │ UNIQUE(user,product) │')
mono('│ id            │                           └──────────────────────┘')
mono('│ user_id (FK)  │')
mono('│ total         │  ┌──────────────────┐     ┌─────────────────────┐')
mono('│ status        │──▶  OrderItem       │     │   UserListing        │')
mono('│ payment_id    │  │──────────────────│     │─────────────────────│')
mono('│ shipping_addr │  │ order_id (FK)    │     │ seller_id (FK→User) │')
mono('│ customer_email│  │ product_id (FK)  │     │ sell_category_id    │')
mono('│ created_at    │  │ quantity         │     │ title               │')
mono('└───────────────┘  │ price_at_purchase│     │ description         │')
mono('                   │ product_name     │     │ asking_price        │')
mono('                   └──────────────────┘     │ condition           │')
mono('                                            │ image_url 1-4       │')
mono('┌──────────────────┐  ┌─────────────────┐  │ status (pending/    │')
mono('│  SellCategory    │──▶  UserListing    │  │  approved/rejected) │')
mono('│──────────────────│  │  (see right)    │  │ admin_note          │')
mono('│ id               │  └─────────────────┘  └─────────────────────┘')
mono('│ name             │')
mono('│ description      │  ┌──────────────────┐  ┌──────────────────────┐')
mono('│ is_active        │  │    Service       │  │     Enquiry          │')
mono('│ sort_order       │  │──────────────────│  │──────────────────────│')
mono('└──────────────────┘  │ id               │  │ id                   │')
mono('                      │ name             │  │ email                │')
mono('                      │ slug (uniq)      │  │ phone                │')
mono('                      │ description      │  │ details              │')
mono('                      │ price            │  │ created_at           │')
mono('                      │ is_active        │  └──────────────────────┘')
mono('                      │ sort_order       │')
mono('                      └──────────────────┘')
doc.add_paragraph()

h2('Model Reference Table')
cols   = ['Model', 'Table', 'Key Fields', 'Relationships']
widths = [1.2, 1.3, 2.8, 2.4]
t = header_table(cols, widths)
models = [
    ('User',         'users',          'id, name, email★, phone, password_hash, is_admin, is_approved',
                                       'Cart 1:1, Orders 1:N, Listings 1:N, Reviews 1:N'),
    ('Category',     'categories',     'id, name, slug★, description',
                                       'Products 1:N'),
    ('Product',      'products',       'id, name, slug★, price, stock, product_type, is_active, is_featured, image_url×4',
                                       'Category N:1, CartItems 1:N, OrderItems 1:N, Reviews 1:N, Likes 1:N'),
    ('Cart',         'carts',          'id, user_id★, created_at',
                                       'User 1:1, CartItems 1:N (cascade delete)'),
    ('CartItem',     'cart_items',     'id, cart_id, product_id, quantity',
                                       'Cart N:1, Product N:1'),
    ('Order',        'orders',         'id, user_id, total, status, payment_intent_id, shipping_address, customer_email',
                                       'User N:1, OrderItems 1:N'),
    ('OrderItem',    'order_items',    'id, order_id, product_id, quantity, price_at_purchase, product_name',
                                       'Order N:1, Product N:1'),
    ('Review',       'reviews',        'id, product_id, user_id, rating(1-5), body | UNIQUE(user, product)',
                                       'Product N:1, User N:1'),
    ('ProductLike',  'product_likes',  'id, user_id, product_id | UNIQUE(user, product)',
                                       'User N:1, Product N:1'),
    ('UserListing',  'user_listings',  'id, seller_id, sell_category_id, title, condition, status, image_url×4, admin_note',
                                       'User(seller) N:1, SellCategory N:1'),
    ('SellCategory', 'sell_categories','id, name, description, is_active, sort_order',
                                       'UserListings 1:N'),
    ('Service',      'services',       'id, name, slug★, description, price, is_active, sort_order',
                                       '(standalone)'),
    ('Enquiry',      'enquiries',      'id, email, phone, details, created_at',
                                       '(standalone — admin read/delete only)'),
]
for i, row in enumerate(models):
    add_row(t, row, widths, alt=(i % 2 == 0))
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 5. REQUEST FLOW DIAGRAMS
# ═════════════════════════════════════════════════════════════════════════════
h1('5.  Request Flow Diagrams')

# ── 5a. Generic request flow ──────────────────────────────────────────────────
h2('5a.  Generic HTTP Request Lifecycle')
mono('Browser')
mono('   │')
mono('   │  GET/POST https://styleus.co.in/some-route')
mono('   ▼')
mono('Nginx  ──── serves /static/ files directly (no Flask involved)')
mono('   │')
mono('   │  proxies dynamic requests to :5000')
mono('   ▼')
mono('Gunicorn  (picks a free worker)')
mono('   │')
mono('   ▼')
mono('Flask app.py  →  matches URL to blueprint route')
mono('   │')
mono('   ├── @login_required?  →  if not logged in → redirect /auth/login')
mono('   │')
mono('   ├── @admin_required?  →  if not admin → 403 Forbidden')
mono('   │')
mono('   ├── Context processors run (cart_count, liked_ids, now)')
mono('   │')
mono('   ├── Route handler executes:')
mono('   │       ├── reads request.form / request.args / request.files')
mono('   │       ├── queries PostgreSQL via SQLAlchemy')
mono('   │       ├── validates data / runs business logic')
mono('   │       └── db.session.commit() if writing data')
mono('   │')
mono('   └── Returns:')
mono('         ├── render_template(...)  →  Jinja2 renders HTML  →  200')
mono('         ├── redirect(url_for(...))  →  302')
mono('         └── jsonify(...)  →  JSON response (cart AJAX)')
doc.add_paragraph()

# ── 5b. Auth flow ─────────────────────────────────────────────────────────────
h2('5b.  Authentication Flow')
mono('                   ┌──────────────────┐')
mono('                   │   New Visitor    │')
mono('                   └────────┬─────────┘')
mono('                            │')
mono('          ┌─────────────────┼──────────────────┐')
mono('          │                 │                  │')
mono('          ▼                 ▼                  ▼')
mono('  ┌──────────────┐  ┌─────────────┐  ┌──────────────────┐')
mono('  │   Register   │  │    Login    │  │  Browse as Guest  │')
mono('  └──────┬───────┘  └──────┬──────┘  │  (session cart)  │')
mono('         │                 │          └──────────────────┘')
mono('         │                 │')
mono('  Name, Email,      Email + Password')
mono('  Phone, Password')
mono('         │                 │')
mono('         ▼                 ▼')
mono('  User saved to DB   Credentials checked')
mono('  is_approved=False       │')
mono('         │          is_approved?')
mono('         │          ├─ NO  → flash warning, stay on login')
mono('         ▼          └─ YES → login_user()')
mono('  Pending approval              │')
mono('  Admin approves via            ▼')
mono('  /admin/users            Merge session cart into DB cart')
mono('         │                      │')
mono('         └──────────────────────┘')
mono('                                │')
mono('                                ▼')
mono('                         Authenticated session')
mono('                         (Flask-Login cookie)')
doc.add_paragraph()

# ── 5c. Cart & Checkout flow ──────────────────────────────────────────────────
h2('5c.  Cart & Checkout Flow')
mono('Product Page')
mono('    │  click "Add to Cart"')
mono('    │  POST /cart/add  (JSON)')
mono('    ▼')
mono('cart.py — add_to_cart()')
mono('    ├── Logged in?')
mono('    │   └── YES → upsert CartItem in PostgreSQL')
mono('    │   └── NO  → store product_id:qty in Flask session dict')
mono('    └── returns { success, cart_count }  → badge updates in JS')
mono('')
mono('Cart Page (/cart/)')
mono('    ├── Quantity +/- → POST /cart/update  (AJAX)')
mono('    ├── Remove       → POST /cart/remove  (AJAX)')
mono('    └── Proceed to Checkout')
mono('              │')
mono('              ▼')
mono('         /checkout/  [currently disabled — flag in checkout.py]')
mono('              │')
mono('              │  POST with shipping details')
mono('              ▼')
mono('         Stripe PaymentIntent.create(amount, currency=inr)')
mono('              │')
mono('              ├─ AuthenticationError (demo keys) → pi_demo_xxx')
mono('              │')
mono('              ▼')
mono('         Order + OrderItems saved to DB')
mono('         Physical product stock decremented')
mono('         Cart cleared (DB or session)')
mono('              │')
mono('              ▼')
mono('         Redirect → /checkout/success')
doc.add_paragraph()

# ── 5d. Used-item listing flow ────────────────────────────────────────────────
h2('5d.  Used-Item Listing Flow  (Sell)')
mono('Logged-in User')
mono('    │  GET /sell')
mono('    ▼')
mono('ListingForm  (title, description, category, price, condition, photos x4)')
mono('    │  POST /sell')
mono('    ▼')
mono('UserListing saved  →  status = "pending"')
mono('    │')
mono('    ▼')
mono('Admin: GET /admin/listings?status=pending')
mono('    │')
mono('    ├── POST /admin/listings/<id>/approve  →  status = "approved"')
mono('    │       → listing appears on /sell/listings (public)')
mono('    │')
mono('    └── POST /admin/listings/<id>/reject   →  status = "rejected"')
mono('            → seller sees "Rejected" on /sell/my-listings')
doc.add_paragraph()

# ── 5e. Tech Enquiry flow ─────────────────────────────────────────────────────
h2('5e.  Tech Enquiry Flow')
mono('Homepage  →  "Tech Enquiries" button')
mono('    │  GET /enquiry')
mono('    ▼')
mono('enquiry.html  form (Email, Phone, Details)')
mono('    │  POST /enquiry')
mono('    ▼')
mono('shop.py — enquiry()')
mono('    ├── Validate email   → regex  ^[a-zA-Z0-9._%+\\-]+@[...]+$')
mono('    ├── Validate phone   → strip spaces/dashes, match ^(\\+91|91|0)?[6-9]\\d{9}$')
mono('    ├── Validate details → len >= 25 characters')
mono('    ├── Any errors? → re-render form with inline error messages')
mono('    └── All valid?  → save Enquiry to DB → flash success → redirect /enquiry')
mono('')
mono('Admin: GET /admin/enquiries')
mono('    └── Table: id, email, phone, details (truncated), date, [Delete button]')
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 6. EXTENSION & MIDDLEWARE STACK
# ═════════════════════════════════════════════════════════════════════════════
h1('6.  Extensions & Middleware Stack')
cols   = ['Extension', 'Purpose', 'Config Key']
widths = [1.8, 4.0, 1.9]
t = header_table(cols, widths)
exts = [
    ('Flask-SQLAlchemy', 'ORM — all DB reads/writes go through SQLAlchemy sessions',          'SQLALCHEMY_DATABASE_URI'),
    ('Flask-Login',      'Session management — login_required, current_user, login_user()',   'login_view = auth.login'),
    ('Flask-WTF / WTForms', 'Form class definitions + CSRF protection on all POST forms',     'WTF_CSRF_ENABLED = True'),
    ('Flask-WTF CSRF',   'Token injected in base.html meta tag; validated on every POST',     'csrf_token()'),
    ('Stripe SDK',       'Payment Intent creation during checkout',                           'STRIPE_SECRET_KEY'),
    ('Werkzeug',         'Password hashing (generate/check_password_hash), file upload utils','—'),
    ('python-slugify',   'Auto-generates URL slugs for products, categories, services',       '—'),
    ('Gunicorn',         'Production WSGI server, 2 sync workers',                           'entrypoint.sh'),
]
for i, row in enumerate(exts):
    add_row(t, row, widths, alt=(i % 2 == 0))
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 7. CONTEXT PROCESSORS (Global template variables)
# ═════════════════════════════════════════════════════════════════════════════
h1('7.  Context Processors  (Global Template Variables)')
body('These functions run on every request and inject variables into every Jinja2 template automatically.')
doc.add_paragraph()
cols   = ['Variable', 'Type', 'Source', 'Used in Templates']
widths = [1.4, 1.0, 2.3, 3.0]
t = header_table(cols, widths)
ctx = [
    ('now',        'datetime', 'datetime.utcnow()',            'Footer copyright year, date formatting'),
    ('cart_count', 'int',      'Cart.get_item_count() or sum(session cart)', 'Navbar cart badge count'),
    ('liked_ids',  'set',      'ProductLike query for current user',          'Product cards — show filled heart icon'),
    ('csrf_token', 'str',      'Flask-WTF CSRFProtect',        'Hidden input in all POST forms'),
    ('current_user','proxy',   'Flask-Login',                  'Nav: show/hide Login, Sell, Admin links'),
]
for i, row in enumerate(ctx):
    add_row(t, row, widths, alt=(i % 2 == 0))
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 8. DOCKER & DEPLOYMENT
# ═════════════════════════════════════════════════════════════════════════════
h1('8.  Docker & Deployment')
h2('Container Setup  (docker-compose.yml at /d/AI/app/)')
cols   = ['Container',   'Image',             'Port',    'Role']
widths = [1.3, 2.9, 0.8, 2.7]
t = header_table(cols, widths)
containers = [
    ('stylus_app', 'registry.aibrainlabs.online/styleus:<tag>', '8082→5000', 'Flask + Gunicorn web server'),
    ('stylus_db',  'postgres:15',                               '(internal)','PostgreSQL database'),
]
for i, row in enumerate(containers):
    add_row(t, row, widths, alt=(i % 2 == 0))

doc.add_paragraph()
h2('Volumes')
bullet('pgdata  →  /var/lib/postgresql/data   (PostgreSQL data persistence)')
bullet('uploads →  /app/static/uploads        (user-uploaded images, survives redeployment)')

doc.add_paragraph()
h2('entrypoint.sh startup sequence')
mono('1.  python -c "from app import app; from models import db;')
mono('       with app.app_context(): db.create_all()"     ← creates missing tables')
mono('2.  flask seed    ← seeds admin user + sample products (skips if already exist)')
mono('3.  gunicorn app:app -w 2 -b 0.0.0.0:5000          ← starts web server')

doc.add_paragraph()
h2('Deployment Workflow')
mono('1.  Edit code in D:\\AI\\styleus\\')
mono('2.  docker build -t registry.aibrainlabs.online/styleus:<YYYYMMDD-HHMMSS> .')
mono('3.  Update image tag in /d/AI/app/docker-compose.yml')
mono('4.  docker compose up -d --force-recreate web')
mono('5.  Ask user before: docker push registry.aibrainlabs.online/styleus:<tag>')

doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 9. ADMIN PANEL SUMMARY
# ═════════════════════════════════════════════════════════════════════════════
h1('9.  Admin Panel Summary')
body('URL: http://localhost:8082/admin   —   requires is_admin=True on the User record.')
body('Layout: responsive sidebar (desktop) / hamburger drawer (mobile). Built on admin/base.html.')
doc.add_paragraph()

cols   = ['Section',        'What admin can do']
widths = [1.8, 5.9]
t = header_table(cols, widths)
admin_sections = [
    ('Dashboard',       'KPI cards (revenue, orders, users, products); recent orders table; pending user & listing alerts'),
    ('Products',        'Create, edit, delete, activate/deactivate products; upload up to 4 images per product'),
    ('Orders',          'List & filter by status; view order detail; update status (pending→processing→shipped→delivered→cancelled)'),
    ('Users',           'List all users; approve pending registrations; grant/revoke admin role'),
    ('Categories',      'Create, edit, delete product categories (unlinking products on delete)'),
    ('User Listings',   'Review pending used-item listings; approve or reject with optional note'),
    ('Sell Categories', 'Manage the categories available when users submit used items for sale'),
    ('Services',        'Create, edit, delete, activate/deactivate IT services shown on /services'),
    ('Enquiries',       'Read-only view of all tech enquiries (email, phone, details, date); delete individual enquiries'),
]
for i, row in enumerate(admin_sections):
    add_row(t, row, widths, alt=(i % 2 == 0))
doc.add_paragraph()

# ═════════════════════════════════════════════════════════════════════════════
# 10. SECURITY NOTES
# ═════════════════════════════════════════════════════════════════════════════
h1('10.  Security Notes')
bullet('CSRF: every POST form includes a hidden csrf_token field validated by Flask-WTF.')
bullet('Passwords: stored as Werkzeug bcrypt hashes — never plain text.')
bullet('Admin protection: @admin_required decorator checks is_admin on current_user; returns 403 otherwise.')
bullet('User approval: new registrations are is_approved=False; login is blocked until admin approves.')
bullet('File uploads: size limited to 2 MB per image (10 MB total); allowed extensions enforced by Flask-WTF FileAllowed.')
bullet('SQL injection: all DB queries go through SQLAlchemy ORM parameterised queries — no raw SQL.')
bullet('Session security: Flask session is cookie-signed with SECRET_KEY from environment variable.')
bullet('Checkout: currently disabled via CHECKOUT_DISABLED flag in checkout.py for safety during development.')

doc.add_paragraph()
rule()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('— End of Document —')
run.font.color.rgb = GREY
run.font.size = Pt(9)

# ── Save ──────────────────────────────────────────────────────────────────────
out = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Styleus_Architecture.docx')
doc.save(out)
print(f'Saved: {out}')
